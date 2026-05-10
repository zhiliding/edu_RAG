from typing import Iterator
from edu_ocr import get_ocr
# 导入必要的模块
from tqdm import tqdm
from docx.table import _Cell, Table  # 用于处理表格
from docx.oxml.table import CT_Tbl  # 用于处理表格XML结构
from docx.oxml.text.paragraph import CT_P  # 用于处理段落XML结构
from docx.text.paragraph import Paragraph  # 用于处理段落内容
from docx import Document as Docu1
from docx.document import Document as Docu2
from docx import ImagePart  # 用于处理Word文档和图片
from PIL import Image  # 用于处理图片
from io import BytesIO  # 用于将字节流转换为图片
import numpy as np  # 用于处理数组
from langchain_core.documents import Document
from langchain_core.document_loaders import BaseLoader


class OCRDOCLoader(BaseLoader):
    """An example document loader that reads a file line by line."""

    def __init__(self, filepath: str) -> None:
        """Initialize the loader with a file path.

        Args:
            filepath_path: The path to the filepath to load.
        """
        self.filepath = filepath

    def lazy_load(self) -> Iterator[Document]:
        # <-- Does not take any arguments
        """A lazy loader that reads a file line by line.

        When you're implementing lazy load methods, you should use a generator
        to yield documents one by one.
        """

        line = self.doc2text(self.filepath)
        yield Document(page_content=line, metadata={"source": self.filepath})

    def doc2text(self, filepath):

        # 创建OCR识别对象
        ocr = get_ocr()
        # print(f'ocr--》{ocr}')  # 输出OCR对象信息

        # 读取Word文档
        doc = Docu1(filepath)
        # print(f'doc-->{doc}')  # 输出读取到的文档信息
        # 定义一个空字符串用于存储最终的文本内容
        resp = ""
        # 定义一个迭代器，用于遍历文档中的块（段落、表格等）
        def iter_block_items(parent):
            # 判断parent对象类型，如果是Document类型，则获取其元素
            if isinstance(parent, Docu2):
                parent_elm = parent.element.body
            # 如果是表格单元格类型，获取单元格的XML元素
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                raise ValueError("OCRDOCLoader parse fail")  # 如果都不是，则抛出错误
            # print(f'parent_elm--》{parent_elm}')
            # print('*'*80)
            # 遍历parent_elm中的所有子元素
            for child in parent_elm.iterchildren():
                # print(f'child--》{child}')
                if isinstance(child, CT_P):  # 如果是段落类型
                    yield Paragraph(child, parent)  # 返回段落
                elif isinstance(child, CT_Tbl):  # 如果是表格类型
                    yield Table(child, parent)  # 返回表格

        # print(f'doc.paragraphs-->{doc.paragraphs}')
        # print(f'doc.tables-->{doc.tables}')
        # 创建进度条，表示文档处理的进度
        b_unit = tqdm(total=len(doc.paragraphs) + len(doc.tables),
                      desc="OCRDOCLoader block index: 0")

        # 遍历文档中的所有块（段落和表格）
        for i, block in enumerate(iter_block_items(doc)):
            # 更新进度条描述
            b_unit.set_description("OCRDOCLoader  block index: {}".format(i))
            b_unit.refresh()  # 刷新进度条

            # 如果块是段落类型
            if isinstance(block, Paragraph):
                resp += block.text.strip() + "\n"  # 将段落文本加入到返回字符串中
                # 获取段落中的所有图片
                images = block._element.xpath('.//pic:pic')
                for image in images:
                    # 遍历图片，获取图片ID
                    for img_id in image.xpath('.//a:blip/@r:embed'):
                        part = doc.part.related_parts[img_id]  # 根据图片ID获取图片对象
                        if isinstance(part, ImagePart):  # 如果该部分是图片
                            # BytesIO 是 Python 内置的 io 模块中的一个类，用于在内存中读写二进制数据
                            # part._blob 通常表示从某个文档（如 DOCX 文件）中提取的二进制内容。
                            image = Image.open(BytesIO(part._blob))  # 打开图片
                            result, _ = ocr(np.array(image))  # 使用OCR识别图片中的文字
                            if result:  # 如果识别结果不为空
                                ocr_result = [line[1] for line in result]  # 提取识别出的文字
                                resp += "\n".join(ocr_result)  # 将识别结果加入返回文本中
            # 如果块是表格类型
            elif isinstance(block, Table):
                # 遍历表格中的所有行和单元格
                for row in block.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            resp += paragraph.text.strip() + "\n"  # 将单元格内的段落文本加入返回文本中

            # 更新进度条
            b_unit.update(1)
        # 返回提取的文本内容
        return resp



if __name__ == '__main__':
    docx_loader = OCRDOCLoader(filepath='/Users/ligang/Desktop/AI29期课堂资料/Codes/integrated_qa_system/rag_qa/samples/ocr_02.docx')
    doc = docx_loader.load()
    print(doc)